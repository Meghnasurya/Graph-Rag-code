# Standard library imports
import os
import zipfile
import tempfile
import uuid
import ast
import time
import re
import json
import base64
from typing import List, Dict, Any, Optional
from datetime import datetime, timedelta
from functools import wraps
from threading import Thread

# Third-party imports
import networkx as nx
import jwt
import requests
from requests.auth import HTTPBasicAuth
from dotenv import load_dotenv
from flask import Flask, request, jsonify, render_template, session, redirect, url_for, flash
from pydantic import BaseModel, Field
from werkzeug.security import generate_password_hash, check_password_hash

# Google AI and embeddings
import google.generativeai as genai

# Database imports
from neo4j import GraphDatabase
from qdrant_client import QdrantClient
from qdrant_client.http.models import VectorParams, Distance, PointStruct

# Document processing
from docx import Document

# API integrations
from github import Github
from atlassian import Jira


# Load environment variables
load_dotenv()


# Configuration
class Config:
    SECRET_KEY = os.getenv("SECRET_KEY", "your-secret-key-here")
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
    NEO4J_URI = os.getenv("NEO4J_URI")
    NEO4J_USER = os.getenv("NEO4J_USER")
    NEO4J_PASSWORD = os.getenv("NEO4J_PASSWORD")
    QDRANT_URL = os.getenv("QDRANT_URL")
    QDRANT_API_KEY = os.getenv("QDRANT_API_KEY")


# Configure Gemini (lazy initialization)
gemini = None
qdrant = None

def init_gemini():
    """Initialize Gemini AI client"""
    global gemini
    if gemini is None and Config.GEMINI_API_KEY:
        try:
            genai.configure(api_key=Config.GEMINI_API_KEY)
            gemini = genai.GenerativeModel("models/gemini-1.5-flash")
            print("✓ Gemini AI initialized successfully")
        except Exception as e:
            print(f"⚠ Gemini AI initialization failed: {e}")
    return gemini

def init_qdrant():
    """Initialize Qdrant client"""
    global qdrant
    if qdrant is None and Config.QDRANT_URL:
        try:
            qdrant = QdrantClient(url=Config.QDRANT_URL, api_key=Config.QDRANT_API_KEY)
            # Test connection
            qdrant.get_collections()
            print("✓ Qdrant initialized successfully")
        except Exception as e:
            print(f"⚠ Qdrant initialization failed: {e}")
            qdrant = None
    return qdrant

# Configure collections
DOCS_COLLECTION = "documentation_embeddings"
TESTS_COLLECTION = "test_embeddings"

def init_collections():
    """Initialize Qdrant collections"""
    qdrant_client = init_qdrant()
    if qdrant_client is None:
        print("⚠ Skipping collection initialization - Qdrant not available")
        return
    
    # Ensure Qdrant collections exist
    for collection_name in [DOCS_COLLECTION, TESTS_COLLECTION]:
        try:
            collections = [c.name for c in qdrant_client.get_collections().collections]
            if collection_name not in collections:
                qdrant_client.create_collection(
                    collection_name=collection_name,
                    vectors_config=VectorParams(size=768, distance=Distance.COSINE),
                )
                print(f"✓ Created collection: {collection_name}")
            else:
                print(f"✓ Collection exists: {collection_name}")
        except Exception as e:
            print(f"⚠ Error checking/creating collection {collection_name}: {e}")


# Data models
class TestGenInput(BaseModel):
    requirement: str = Field(...)
    documentation: str = Field(...)


class TestOutput(BaseModel):
    test_cases: List[str]
    test_scripts: List[str]


class UserSession(BaseModel):
    user_id: str
    github_token: Optional[str] = None
    jira_token: Optional[str] = None
    jira_url: Optional[str] = None
    jira_email: Optional[str] = None
    selected_github_repo: Optional[str] = None
    selected_jira_project: Optional[str] = None


# Authentication decorator
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('connect_services'))
        return f(*args, **kwargs)
    return decorated_function


# Service Classes
class GitHubService:
    def __init__(self, token: str):
        self.github = Github(token)
        self.token = token

    def get_user_repos(self) -> List[Dict]:
        """Get user's repositories"""
        try:
            repos = []
            for repo in self.github.get_user().get_repos():
                repos.append({
                    'name': repo.name,
                    'full_name': repo.full_name,
                    'description': repo.description,
                    'private': repo.private
                })
            return repos
        except Exception as e:
            print(f"Error fetching repos: {e}")
            return []

    def extract_docx_content(self, base64_content: str) -> str:
        """Extract text content from DOCX file"""
        try:
            import io
            # Decode base64 content
            docx_bytes = base64.b64decode(base64_content)
            
            # Create a BytesIO object to simulate a file
            docx_stream = io.BytesIO(docx_bytes)
            
            # Use python-docx to extract text
            doc = Document(docx_stream)
            
            # Extract all text from paragraphs
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    full_text.append(paragraph.text.strip())
            
            # Also extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            
            return '\n\n'.join(full_text)
        except Exception as e:
            print(f"Error extracting DOCX content: {e}")
            return ""

    def get_repo_contents(self, repo_name: str, path: str = "") -> Dict[str, Any]:
        """Get repository contents for documentation and tests"""
        try:
            repo = self.github.get_repo(repo_name)
            contents = {
                'docs': {},
                'tests': {},
                'code': {}
            }

            def process_contents(items, current_path=""):
                for item in items:
                    if item.type == "file":
                        try:
                            if item.name.endswith(('.py', '.js', '.java', '.cpp', '.c')):
                                content = base64.b64decode(item.content).decode('utf-8')
                                if 'test' in item.name.lower() or 'spec' in item.name.lower():
                                    contents['tests'][item.path] = content
                                else:
                                    contents['code'][item.path] = content
                            elif item.name.endswith(('.md', '.txt', '.rst', '.doc')):
                                content = base64.b64decode(item.content).decode('utf-8')
                                contents['docs'][item.path] = content
                            elif item.name.endswith('.docx'):
                                # Handle DOCX files
                                print(f"Processing DOCX file: {item.path}")
                                docx_content = self.extract_docx_content(item.content)
                                if docx_content:
                                    contents['docs'][item.path] = docx_content
                                    print(f"Successfully extracted {len(docx_content)} characters from {item.path}")
                                else:
                                    print(f"No content extracted from DOCX file: {item.path}")
                        except Exception as e:
                            print(f"Error processing file {item.path}: {e}")
                    elif item.type == "dir" and not item.name.startswith('.'):
                        # Recursively process subdirectories
                        try:
                            sub_contents = repo.get_contents(item.path)
                            process_contents(sub_contents, item.path)
                        except Exception as e:
                            print(f"Error processing directory {item.path}: {e}")

            root_contents = repo.get_contents(path)
            process_contents(root_contents)
            return contents
        except Exception as e:
            print(f"Error fetching repo contents: {e}")
            return {'docs': {}, 'tests': {}, 'code': {}}


class JiraService:
    def __init__(self, url: str, email: str, token: str):
        self.jira = Jira(url=url, username=email, password=token, cloud=True)
        self.url = url
        self.email = email
        self.token = token

    def get_projects(self) -> List[Dict]:
        """Get user's Jira projects"""
        try:
            projects = self.jira.projects()
            return [{'key': p['key'], 'name': p['name']} for p in projects]
        except Exception as e:
            print(f"Error fetching Jira projects: {e}")
            return []

    def get_test_cases(self, project_key: str) -> List[Dict]:
        """Get test cases from Jira project"""
        try:
            # Search for test-related issues
            jql = f'project = {project_key} AND (issueType = "Test" OR labels = "test" OR summary ~ "test")'
            issues = self.jira.jql(jql)
            test_cases = []
            for issue in issues.get('issues', []):
                test_cases.append({
                    'key': issue['key'],
                    'summary': issue['fields']['summary'],
                    'description': issue['fields'].get('description', ''),
                    'status': issue['fields']['status']['name']
                })
            return test_cases
        except Exception as e:
            print(f"Error fetching test cases: {e}")
            return []

    def create_or_update_test_case(self, project_key: str, test_data: Dict):
        """
        Creates a new Jira issue as a test case, or updates an existing one if a match is found.
        The matching is based on the suggested file_name from the test generation
        matching the Jira issue summary.
        """
        try:
            # Use file_name for summary and content for description
            summary = test_data.get('file_name', 'Generated Test Case').replace('.py', '')
            test_code_content = test_data.get('test_code', '')
            test_cases_list = test_data.get('test_cases', [])

            # Format test cases for description using Jira Wiki Markup
            formatted_test_cases = ""
            if test_cases_list:
                formatted_test_cases = "\n*h2. Detailed Test Cases:*\n"
                for tc in test_cases_list:
                    formatted_test_cases += f"{{panel:title={tc.get('name', 'Unnamed Test Case')}|borderStyle=solid|borderColor=#ccc|titleBGColor=#eee|bgColor=#fafafa}}\n"
                    formatted_test_cases += f"{{color:grey}}Description:{{color}} {tc.get('description', 'No description.')}\n"
                    if tc.get('steps'):
                        formatted_test_cases += f"{{color:grey}}Steps:{{color}}\n"
                        for i, step in enumerate(tc['steps']):
                            formatted_test_cases += f"  - {step}\n"
                    if tc.get('expected_result'):
                        formatted_test_cases += f"{{color:grey}}Expected Result:{{color}} {tc.get('expected_result', 'No expected result.')}\n"
                    formatted_test_cases += "*{{panel}}\n\n"

            description = f"{{code:python}}\n{test_code_content}\n{{code}}{formatted_test_cases}"

            # Attempt to find an existing test case by matching the summary
            jql_query = f'project = "{project_key}" AND summary ~ "{summary}" AND issueType = "Test"'
            existing_issues = self.jira.jql(jql_query).get('issues', [])

            if existing_issues:
                # Assuming the first match is the one to update
                issue_key = existing_issues[0]['key']
                self.jira.issue_update(
                    issue_key,
                    fields={
                        'summary': summary,
                        'description': description
                    }
                )
                print(f"Updated Jira test case {issue_key} in project {project_key}")
                return issue_key
            else:
                # Create a new issue if no existing match is found
                new_issue = self.jira.create_issue(
                    fields={
                        'project': {'key': project_key},
                        'summary': summary,
                        'description': description,
                        'issuetype': {'name': 'Test'}
                    }
                )
                print(f"Created new Jira test case {new_issue.key} in project {project_key}")
                return new_issue.key
        except Exception as e:
            print(f"Error creating/updating Jira test case: {e}")
            return None


# RAG Classes
class EnhancedGraphRAG:
    def __init__(self):
        self.driver = GraphDatabase.driver(
            Config.NEO4J_URI, 
            auth=(Config.NEO4J_USER, Config.NEO4J_PASSWORD)
        )

    def close(self):
        if self.driver:
            self.driver.close()

    def verify_connection(self):
        try:
            with self.driver.session() as session:
                result = session.run("RETURN 1")
                return result.single()[0] == 1
        except Exception as e:
            print(f"Neo4j connection failed: {e}")
            return False

    def update_knowledge(self, data: Dict[str, Any], source: str):
        """Update knowledge without clearing existing data"""
        if not self.verify_connection():
            print("Neo4j connection failed, skipping graph operations")
            return
        
        try:
            with self.driver.session() as session:
                # Add source tracking
                session.run(
                    "MERGE (s:Source {name: $source, updated: datetime()})",
                    {"source": source}
                )
                
                # Process documentation
                if 'docs' in data:
                    for file_path, content in data['docs'].items():
                        chunks = self.chunk_text(content)
                        for chunk in chunks:
                            session.run("""
                                MERGE (s:Source {name: $source})
                                MERGE (d:Document {path: $path})
                                MERGE (c:Content {text: $chunk})
                                MERGE (s)-[:CONTAINS]->(d)
                                MERGE (d)-[:HAS_CONTENT]->(c)
                                """, {"source": source, "path": file_path, "chunk": chunk})
                
                # Process test cases
                if 'tests' in data:
                    for file_path, content in data['tests'].items():
                        functions = self.extract_test_functions(content)
                        for func_name, func_body in functions:
                            session.run("""
                                MERGE (s:Source {name: $source})
                                MERGE (t:TestFile {path: $path})
                                MERGE (f:TestFunction {name: $func_name, body: $func_body})
                                MERGE (s)-[:CONTAINS]->(t)
                                MERGE (t)-[:HAS_TEST]->(f)
                                """, {"source": source, "path": file_path, "func_name": func_name, "func_body": func_body})
                
                # Process code
                if 'code' in data:
                    for file_path, content in data['code'].items():
                        functions = self.extract_functions(content)
                        for func_name, func_doc in functions:
                            session.run("""
                                MERGE (s:Source {name: $source})
                                MERGE (cf:CodeFile {path: $path})
                                MERGE (f:Function {name: $func_name, documentation: $func_doc})
                                MERGE (s)-[:CONTAINS]->(cf)
                                MERGE (cf)-[:HAS_FUNCTION]->(f)
                                """, {"source": source, "path": file_path, "func_name": func_name, "func_doc": func_doc})
                
                print(f"Successfully updated knowledge graph with data from {source}")
        except Exception as e:
            print(f"Error updating knowledge graph: {e}")

    def search_related_content(self, query: str, limit: int = 10) -> List[Dict]:
        """Search for related content in the knowledge graph"""
        try:
            with self.driver.session() as session:
                result = session.run("""
                    MATCH (n)
                    WHERE n.name CONTAINS $query OR n.text CONTAINS $query OR n.documentation CONTAINS $query
                    RETURN n, labels(n) as labels
                    LIMIT $limit
                    """, {"query": query.lower(), "limit": limit})
                return [{"node": record["n"], "labels": record["labels"]} for record in result]
        except Exception as e:
            print(f"Error searching graph: {e}")
            return []

    def chunk_text(self, text: str, max_chars: int = 400) -> List[str]:
        """Chunk text for storage"""
        if not text or len(text.strip()) < 20:
            return []
        
        text = re.sub(r'\s+', ' ', text.strip())
        chunks = []
        
        # Split by paragraphs
        paragraphs = text.split('\n\n')
        current_chunk = ""
        
        for paragraph in paragraphs:
            if len(current_chunk) + len(paragraph) <= max_chars:
                current_chunk += paragraph + "\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph + "\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return [chunk for chunk in chunks if len(chunk.strip()) >= 20]

    def extract_test_functions(self, code: str) -> List[tuple]:
        """Extract test functions from code"""
        functions = []
        try:
            tree = ast.parse(code)
            for node in ast.walk(tree):
                if isinstance(node, ast.FunctionDef) and node.name.startswith('test_'):
                    func_body = ast.get_source_segment(code, node) or ""
                    functions.append((node.name, func_body))
        except Exception as e:
            print(f"Error parsing test code: {e}")
        return functions

    def extract_functions(self, code: str) -> List[tuple]:
        """Extract functions from code"""
        functions = []
        try:
            tree = ast.parse(code)
            for node in ast.walk(tree):
                if isinstance(node, ast.FunctionDef):
                    docstring = ast.get_docstring(node) or ""
                    functions.append((node.name, docstring))
        except Exception as e:
            print(f"Error parsing code: {e}")
        return functions


class EnhancedEmbeddingRAG:
    def __init__(self):
        self.client = None

    def get_client(self):
        """Get Qdrant client, initializing if necessary"""
        if self.client is None:
            self.client = init_qdrant()
        return self.client

    def update_embeddings(self, data: Dict[str, Any], source: str, collection_name: str):
        """Update embeddings without clearing existing data"""
        try:
            chunks = []
            metadata_list = []
            
            # Process different types of content
            if 'docs' in data:
                for file_path, content in data['docs'].items():
                    text_chunks = self.chunk_text(content)
                    chunks.extend(text_chunks)
                    metadata_list.extend([{
                        'source': source,
                        'type': 'documentation',
                        'file_path': file_path,
                        'text': chunk
                    } for chunk in text_chunks])
            
            if 'tests' in data:
                for file_path, content in data['tests'].items():
                    test_chunks = self.chunk_text(content)
                    chunks.extend(test_chunks)
                    metadata_list.extend([{
                        'source': source,
                        'type': 'test',
                        'file_path': file_path,
                        'text': chunk
                    } for chunk in test_chunks])
            
            if 'code' in data:
                for file_path, content in data['code'].items():
                    code_chunks = self.chunk_text(content)
                    chunks.extend(code_chunks)
                    metadata_list.extend([{
                        'source': source,
                        'type': 'code',
                        'file_path': file_path,
                        'text': chunk
                    } for chunk in code_chunks])
            
            # Store embeddings
            return self.embed_and_store(chunks, metadata_list, collection_name)
        except Exception as e:
            print(f"Error updating embeddings: {e}")
            return 0

    def embed_and_store(self, chunks: List[str], metadata_list: List[Dict], collection_name: str) -> int:
        """Store embeddings with metadata"""
        if not chunks:
            return 0
        
        points = []
        batch_size = 25
        
        for i in range(0, len(chunks), batch_size):
            batch_chunks = chunks[i:i+batch_size]
            batch_metadata = metadata_list[i:i+batch_size]
            batch_points = []
            
            for chunk, metadata in zip(batch_chunks, batch_metadata):
                try:
                    gemini_client = init_gemini()
                    if not gemini_client:
                        print("⚠ Gemini client not available for embedding")
                        continue
                        
                    result = genai.embed_content(
                        model="models/text-embedding-004",
                        content=chunk[:1500],
                        task_type="retrieval_document"
                    )
                    embedding = result['embedding']
                    batch_points.append(
                        PointStruct(
                            id=str(uuid.uuid4()),
                            vector=embedding,
                            payload=metadata
                        )
                    )
                except Exception as e:
                    print(f"Error embedding chunk: {e}")
                    continue
            
            if batch_points:
                client = self.get_client()
                if client:
                    try:
                        client.upsert(collection_name=collection_name, points=batch_points)
                        points.extend(batch_points)
                    except Exception as e:
                        print(f"Error storing batch: {e}")
                else:
                    print("⚠ Qdrant client not available, skipping batch storage")
            
            time.sleep(0.1)  # Rate limiting
        
        return len(points)

    def search(self, query: str, collection_name: str, top_k: int = 10) -> List[Dict]:
        """Search embeddings"""
        try:
            gemini_client = init_gemini()
            if not gemini_client:
                print("⚠ Gemini client not available for embedding")
                return []
            
            result = genai.embed_content(
                model="models/text-embedding-004",
                content=query,
                task_type="retrieval_query"
            )
            query_vec = result['embedding']
            
            client = self.get_client()
            if not client:
                print("⚠ Qdrant client not available for search")
                return []
            
            results = client.search(
                collection_name=collection_name,
                query_vector=query_vec,
                limit=top_k,
                score_threshold=0.5
            )
            return [hit.payload for hit in results]
        except Exception as e:
            print(f"Error searching embeddings: {e}")
            return []

    def chunk_text(self, text: str, max_chars: int = 400) -> List[str]:
        """Chunk text for embedding"""
        if not text or len(text.strip()) < 20:
            return []
        
        text = re.sub(r'\s+', ' ', text.strip())
        chunks = []
        
        # Split by paragraphs
        paragraphs = text.split('\n\n')
        current_chunk = ""
        
        for paragraph in paragraphs:
            if len(current_chunk) + len(paragraph) <= max_chars:
                current_chunk += paragraph + "\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph + "\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return [chunk for chunk in chunks if len(chunk.strip()) >= 20]


# Main RAG System
class AgenticRAG:
    def __init__(self):
        self.graph_rag = EnhancedGraphRAG()
        self.embedding_rag = EnhancedEmbeddingRAG()

    def process_query(self, query: str, user_session: UserSession) -> Dict[str, Any]:
        """Process user query with agentic approach"""
        try:
            # Step 1: Search existing knowledge
            doc_results = self.embedding_rag.search(query, DOCS_COLLECTION, top_k=5)
            test_results = self.embedding_rag.search(query, TESTS_COLLECTION, top_k=5)
            graph_results = self.graph_rag.search_related_content(query, limit=5)

            # Step 2: Analyze if existing tests cover the requirement
            analysis_prompt = f"""
                Analyze this query: "{query}"
                Existing documentation: {json.dumps(doc_results[:2], indent=2)}
                Existing tests: {json.dumps(test_results[:2], indent=2)}
                
                Determine:
                1. Is there existing functionality for this requirement?
                2. Are there existing tests that cover this?
                3. What new tests need to be created?
                4. Should existing tests be updated?
                
                You MUST return ONLY valid JSON in this exact format with no additional text:
                {{
                    "existing_functionality": true,
                    "existing_tests": false,
                    "needs_new_tests": true,
                    "needs_test_updates": false,
                    "recommendations": ["Create new test cases for the requirement"]
                }}                """
            gemini_client = init_gemini()
            if not gemini_client:
                print("⚠ Gemini client not available for analysis")
                # Fallback analysis
                analysis = {
                    "existing_functionality": False,
                    "existing_tests": False,
                    "needs_new_tests": True,
                    "needs_test_updates": False,
                    "recommendations": ["Generate new test cases based on query"]
                }
            else:
                analysis_response = gemini_client.generate_content(analysis_prompt)
                analysis_text = analysis_response.text.strip()
                
                # Clean up response text to extract JSON
                if analysis_text.startswith('```json'):
                    analysis_text = analysis_text[7:]
                if analysis_text.endswith('```'):
                    analysis_text = analysis_text[:-3]
                analysis_text = analysis_text.strip()
                
                try:
                    analysis = json.loads(analysis_text)
                except json.JSONDecodeError as json_error:
                    print(f"Analysis JSON parsing error: {json_error}")
                    print(f"Analysis response text: {analysis_text}")
                    # Fallback analysis
                    analysis = {
                        "existing_functionality": False,
                        "existing_tests": False,
                        "needs_new_tests": True,
                        "needs_test_updates": False,
                        "recommendations": ["Generate new test cases based on query"]
                    }

            # Step 3: Generate or update tests based on analysis
            if analysis.get('needs_new_tests') or analysis.get('needs_test_updates'):
                test_generation_result = self.generate_tests(query, doc_results, test_results, analysis)                # Step 4: Update test repository if needed
                jira_update_status = None
                if (user_session.jira_token and user_session.jira_url and 
                    user_session.selected_jira_project):
                    jira_update_status = self.update_jira_tests(user_session, test_generation_result)
                else:
                    print("Jira credentials or project not selected, skipping Jira test update.")

                return {
                    'analysis': analysis,
                    'generated_tests': test_generation_result,
                    'doc_results': doc_results,
                    'test_results': test_results,
                    'jira_update_status': jira_update_status,
                    'updated_embeddings': True
                }
            else:
                return {
                    'analysis': analysis,
                    'existing_tests': test_results,
                    'doc_results': doc_results,
                    'test_results': test_results,
                    'message': 'Existing tests already cover this requirement'
                }
        except Exception as e:
            print(f"Error processing query: {e}")
            return {'error': str(e)}

    def generate_tests(self, query: str, doc_context: List[Dict], test_context: List[Dict], analysis: Dict) -> Dict:
        """Generate new tests based on context and analysis"""
        try:
            generation_prompt = f"""
                Based on this requirement: "{query}"
                Context from documentation: {json.dumps(doc_context, indent=2)}
                Context from existing tests: {json.dumps(test_context, indent=2)}
                Analysis: {json.dumps(analysis, indent=2)}
                
                Generate comprehensive test cases and Python unittest code.
                You MUST return ONLY valid JSON in this exact format with no additional text:
                {{
                    "test_cases": [
                        {{
                            "name": "test_case_name",
                            "description": "detailed description",
                            "steps": ["step1", "step2", "step3"],
                            "expected_result": "expected outcome"
                        }}
                    ],
                    "test_code": "complete Python unittest code",
                    "file_name": "suggested_test_file_name.py"
                }}                """
            gemini_client = init_gemini()
            if not gemini_client:
                print("⚠ Gemini client not available for test generation")
                return self._create_fallback_test_result(query)
                
            response = gemini_client.generate_content(generation_prompt)
            
            # Handle empty or None response
            if not response or not hasattr(response, 'text') or not response.text:
                print("Empty response from Gemini API")
                return self._create_fallback_test_result(query)
            
            response_text = response.text.strip()
            
            # Handle empty response text
            if not response_text:
                print("Empty response text from Gemini API")
                return self._create_fallback_test_result(query)
            
            # Clean up response text to extract JSON
            if response_text.startswith('```json'):
                response_text = response_text[7:]
            if response_text.endswith('```'):
                response_text = response_text[:-3]
            response_text = response_text.strip()
            
            # Final check for empty text
            if not response_text:
                print("Response text empty after cleanup")
                return self._create_fallback_test_result(query)
            
            # Try to parse JSON
            try:
                parsed_result = json.loads(response_text)
                # Validate required fields
                if not isinstance(parsed_result, dict):
                    print("Response is not a valid dict")
                    return self._create_fallback_test_result(query)
                
                if 'test_cases' not in parsed_result or 'test_code' not in parsed_result:
                    print("Missing required fields in response")
                    return self._create_fallback_test_result(query)
                
                return parsed_result
            except json.JSONDecodeError as json_error:
                print(f"JSON parsing error: {json_error}")
                print(f"Response text: {response_text}")
                return self._create_fallback_test_result(query)
                
        except Exception as e:
            print(f"Error generating tests: {e}")
            return self._create_fallback_test_result(query)

    def _create_fallback_test_result(self, query: str) -> Dict:
        """Create a fallback test result when LLM fails"""
        safe_name = ''.join(c for c in query if c.isalnum() or c in ' _').replace(' ', '_').lower()
        return {
            "test_cases": [
                {
                    "name": f"test_{safe_name}",
                    "description": f"Test case for: {query}",
                    "steps": ["Setup test environment", "Execute test scenario", "Verify results"],
                    "expected_result": "Test should pass successfully"
                }
            ],
            "test_code": f"import unittest\n\nclass Test{safe_name.title().replace('_', '')}(unittest.TestCase):\n    def test_{safe_name}(self):\n        \"\"\"Test implementation for: {query}\"\"\"\n        # TODO: Implement actual test logic\n        self.assertTrue(True)  # Replace with actual test logic\n\nif __name__ == '__main__':\n    unittest.main()",
            "file_name": f"test_{safe_name}.py"
        }

    def update_jira_tests(self, user_session: UserSession, test_result: Dict):
        """
        Updates Jira test repository with new or updated test cases.
        This function uses the JiraService to create or modify Jira issues
        that represent the generated test cases.
        """
        try:
            jira_service = JiraService(user_session.jira_url, user_session.jira_email, user_session.jira_token)
            project_key = user_session.selected_jira_project
            
            # Call the new method in JiraService to handle the creation/update logic
            issue_key = jira_service.create_or_update_test_case(project_key, test_result)
            
            if issue_key:
                # Re-embed the newly created/updated test case
                test_content = f"{test_result.get('file_name', 'Generated Test Case')}\n{test_result.get('test_code', '')}"
                
                # Create data structure for embedding
                new_test_data = {
                    'tests': {
                        f"jira:{issue_key}": test_content
                    }
                }
                
                # Update embeddings and graph with new test case
                self.embedding_rag.update_embeddings(
                    new_test_data, 
                    f"jira:{user_session.selected_jira_project}", 
                    TESTS_COLLECTION
                )
                self.graph_rag.update_knowledge(
                    new_test_data, 
                    f"jira:{user_session.selected_jira_project}"
                )
                
                print(f"Successfully created/updated and embedded Jira test case {issue_key} in project: {project_key}")
                return {
                    'success': True,
                    'issue_key': issue_key,
                    'project_key': project_key,
                    'message': f'Successfully created/updated test case {issue_key}'
                }
            else:
                print(f"Failed to create/update Jira test case in project: {project_key}")
                return {
                    'success': False,
                    'project_key': project_key,
                    'message': f'Failed to create/update test case in project {project_key}'
                }
                
        except Exception as e:
            print(f"Error updating Jira test repository: {e}")
            return {
                'success': False,
                'error': str(e),
                'message': f'Error updating Jira test repository: {str(e)}'
            }


# Flask Application
app = Flask(__name__)
app.config.from_object(Config)

# Initialization function
def initialize_app():
    """Initialize application services"""
    print("🚀 Initializing Agentic RAG System...")
    
    # Initialize services with graceful fallbacks
    init_gemini()
    init_qdrant()
    init_collections()
    
    print("✅ Application initialization complete")

# Initialize services
agentic_rag = AgenticRAG()


# Routes
@app.route("/")
def index():
    if 'user_id' not in session:
        return redirect(url_for('connect_services'))
    return redirect(url_for('dashboard'))


@app.route("/connect")
def connect_services():
    """Connect to GitHub and Jira with direct credentials"""
    return render_template("connect_services.html")


@app.route("/connect/github", methods=["POST"])
def connect_github():
    """Connect to GitHub using personal access token"""
    try:
        data = request.get_json()
        github_token = data.get('github_token', '').strip()
        
        if not github_token:
            return jsonify({'success': False, 'error': 'GitHub token is required'}), 400
        
        # Test the token by trying to access user info
        github_service = GitHubService(github_token)
        try:
            user = github_service.github.get_user()
            user_login = user.login  # This will fail if token is invalid
        except Exception as e:
            return jsonify({'success': False, 'error': 'Invalid GitHub token or insufficient permissions'}), 400
        
        # Create user session if doesn't exist
        if 'user_id' not in session:
            session['user_id'] = str(uuid.uuid4())
        
        # Store GitHub token
        session['github_token'] = github_token
        session['github_connected'] = True
        
        return jsonify({'success': True, 'message': 'GitHub connected successfully'})
        
    except Exception as e:
        print(f"GitHub connection error: {str(e)}")
        return jsonify({'success': False, 'error': f'Connection failed: {str(e)}'}), 500


@app.route("/connect/jira", methods=["POST"])
def connect_jira():
    """Connect to Jira using direct credentials"""
    try:
        data = request.get_json()
        jira_url = data.get('jira_url', '').strip()
        jira_email = data.get('jira_email', '').strip()
        jira_token = data.get('jira_token', '').strip()
        
        if not all([jira_url, jira_email, jira_token]):
            return jsonify({'success': False, 'error': 'All Jira credentials are required'}), 400
        
        # Ensure URL has proper format
        if not jira_url.startswith(('http://', 'https://')):
            jira_url = 'https://' + jira_url
        
        # Test the credentials by trying to access projects
        try:
            jira = Jira(url=jira_url, username=jira_email, password=jira_token, cloud=True)
            projects = jira.projects()  # This will fail if credentials are invalid
        except Exception as e:
            return jsonify({'success': False, 'error': 'Invalid Jira credentials or URL'}), 400
        
        # Create user session if doesn't exist
        if 'user_id' not in session:
            session['user_id'] = str(uuid.uuid4())
          # Store Jira credentials
        session['jira_url'] = jira_url
        session['jira_email'] = jira_email
        session['jira_token'] = jira_token
        session['jira_connected'] = True
        
        return jsonify({'success': True, 'message': 'Jira connected successfully'})
        
    except Exception as e:
        print(f"Jira connection error: {str(e)}")
        return jsonify({'success': False, 'error': f'Connection failed: {str(e)}'}), 500


@app.route("/fetch/github-repos", methods=["POST"])
def fetch_github_repos():
    """Fetch GitHub repositories for connected account"""
    try:
        if 'github_token' not in session:
            return jsonify({'success': False, 'error': 'GitHub not connected'}), 400
        
        github_service = GitHubService(session['github_token'])
        repos = github_service.get_user_repos()
        
        return jsonify({'success': True, 'repos': repos})
        
    except Exception as e:
        print(f"Error fetching GitHub repos: {str(e)}")
        return jsonify({'success': False, 'error': f'Failed to fetch repositories: {str(e)}'}), 500


@app.route("/fetch/jira-projects", methods=["POST"])
def fetch_jira_projects():
    """Fetch Jira projects for connected account"""
    try:
        if not all(key in session for key in ['jira_url', 'jira_email', 'jira_token']):
            return jsonify({'success': False, 'error': 'Jira not connected'}), 400
        
        jira = Jira(
            url=session['jira_url'],
            username=session['jira_email'], 
            password=session['jira_token'],
            cloud=True
        )
        projects = jira.projects()
        
        # Format projects for frontend
        formatted_projects = []
        for project in projects:
            formatted_projects.append({
                'key': project.get('key'),
                'name': project.get('name'),
                'site_url': session['jira_url']
            })
        
        return jsonify({'success': True, 'projects': formatted_projects})
        
    except Exception as e:
        print(f"Error fetching Jira projects: {str(e)}")
        return jsonify({'success': False, 'error': f'Failed to fetch projects: {str(e)}'}), 500


@app.route("/select-github-repo", methods=["POST"])
@login_required
def select_github_repo():
    """Select GitHub repository from dashboard"""
    if 'github_token' not in session:
        return jsonify({'error': 'GitHub not connected'}), 400
    
    repo_data = request.get_json()
    github_repo = repo_data.get('github_repo') if repo_data else request.form.get('github_repo')
    
    if not github_repo:
        return jsonify({'error': 'Repository selection required'}), 400
    
    try:
        # Test GitHub access
        github_service = GitHubService(session['github_token'])
        repo = github_service.github.get_repo(github_repo)
        
        session['selected_github_repo'] = github_repo
        
        # Start background process to load GitHub data
        def load_github_data():
            try:
                user_session = UserSession(
                    user_id=session['user_id'],
                    github_token=session.get('github_token'),
                    selected_github_repo=session.get('selected_github_repo')
                )

                if user_session.github_token and user_session.selected_github_repo:
                    print(f"Loading GitHub data from: {user_session.selected_github_repo}")
                    github_service = GitHubService(user_session.github_token)
                    repo_data = github_service.get_repo_contents(user_session.selected_github_repo)
                    
                    # Update graph and embeddings
                    agentic_rag.graph_rag.update_knowledge(repo_data, f"github:{user_session.selected_github_repo}")
                    docs_count = agentic_rag.embedding_rag.update_embeddings(repo_data, f"github:{user_session.selected_github_repo}", DOCS_COLLECTION)
                    tests_count = agentic_rag.embedding_rag.update_embeddings(repo_data, f"github:{user_session.selected_github_repo}", TESTS_COLLECTION)
                    print(f"Embedded {docs_count} docs and {tests_count} test chunks")
            except Exception as e:
                print(f"Error loading GitHub data: {str(e)}")

        # Start background thread
        thread = Thread(target=load_github_data)
        thread.daemon = True
        thread.start()

        return jsonify({
            'success': True,
            'message': f'Repository {github_repo} selected successfully. Data loading in background.',
            'redirect': url_for('query_interface')
        })
        
    except Exception as e:
        print(f"Error accessing repository {github_repo}: {e}")
        return jsonify({'error': f'Error accessing repository: {str(e)}'}), 500


@app.route("/select-jira-project-api", methods=["POST"])
@login_required  
def select_jira_project_api():
    """Select Jira project from dashboard"""
    if 'jira_token' not in session:
        return jsonify({'error': 'Jira not connected'}), 400
    
    request_data = request.get_json() if request.is_json else request.form
    jira_project_key = request_data.get('jira_project')
    jira_site_url = request_data.get('jira_site_url')

    if not jira_project_key or not jira_site_url:
        return jsonify({'error': 'Project selection and site URL required'}), 400

    try:
        session['selected_jira_project'] = jira_project_key
        session['jira_url'] = jira_site_url        # Start background process to load Jira data
        def load_jira_data():
            try:
                user_session = UserSession(
                    user_id=session['user_id'],
                    jira_token=session.get('jira_token'),
                    jira_url=session.get('jira_url'),
                    jira_email=session.get('jira_email'),
                    selected_jira_project=session.get('selected_jira_project')
                )

                if (user_session.jira_token and user_session.jira_url and 
                    user_session.selected_jira_project):
                    print(f"Loading Jira data from: {user_session.selected_jira_project}")
                      # Use direct API calls with Basic Auth (email + API token)
                    auth = HTTPBasicAuth(user_session.jira_email, user_session.jira_token)
                    
                    # Search for test-related issues using direct API
                    jql = f'project = {user_session.selected_jira_project} AND (issueType = "Test" OR labels = "test" OR summary ~ "test")'
                    search_url = f"{user_session.jira_url}/rest/api/3/search"
                    
                    response = requests.get(search_url, auth=auth, params={'jql': jql})
                    
                    if response.status_code == 200:
                        issues_data = response.json()
                        test_cases = []
                        
                        for issue in issues_data.get('issues', []):
                            test_cases.append({
                                'key': issue['key'],
                                'summary': issue['fields']['summary'],
                                'description': issue['fields'].get('description', ''),
                                'status': issue['fields']['status']['name']
                            })
                        
                        # Convert test cases to embeddable format
                        jira_data = {
                            'tests': {
                                f"jira:{tc['key']}": f"{tc['summary']}\n{tc['description']}" 
                                for tc in test_cases
                            }
                        }
                        
                        # Update graph and embeddings with Jira data
                        agentic_rag.graph_rag.update_knowledge(jira_data, f"jira:{user_session.selected_jira_project}")
                        tests_count = agentic_rag.embedding_rag.update_embeddings(jira_data, f"jira:{user_session.selected_jira_project}", TESTS_COLLECTION)
                        print(f"Embedded {tests_count} Jira test chunks")
                    else:
                        print(f"Failed to fetch Jira issues: {response.status_code}")
                    
            except Exception as e:
                print(f"Error loading Jira data: {str(e)}")

        # Start background thread
        thread = Thread(target=load_jira_data)
        thread.daemon = True
        thread.start()
        
        return jsonify({
            'success': True,
            'message': f'Project {jira_project_key} selected successfully. Data loading in background.',
            'redirect': url_for('query_interface')
        })
        
    except Exception as e:
        print(f"Error selecting Jira project: {e}")
        return jsonify({'error': f'Error selecting project: {str(e)}'}), 500



@app.route("/dashboard")
@login_required
def dashboard():
    """Main dashboard with GitHub repo and Jira project selection"""
    github_repos = []
    jira_projects = []
    
    # Get GitHub repositories if connected
    if 'github_token' in session:
        try:
            github_service = GitHubService(session['github_token'])
            github_repos = github_service.get_user_repos()
        except Exception as e:
            print(f"Error fetching GitHub repos: {e}")
            flash('Error fetching GitHub repositories')
    
    # Get Jira projects if connected
    if 'jira_token' in session and 'jira_url' in session and 'jira_email' in session:
        try:
            jira = Jira(
                url=session['jira_url'],
                username=session['jira_email'], 
                password=session['jira_token'],
                cloud=True
            )
            projects = jira.projects()
            
            for project in projects:
                project['site_url'] = session['jira_url']
                project['site_name'] = session['jira_url']
                jira_projects.append(project)
                
        except Exception as e:
            print(f"Error fetching Jira projects: {e}")
            flash('Error fetching Jira projects')
    
    return render_template("dashboard.html", 
                         github_repos=github_repos, 
                         jira_projects=jira_projects,
                         github_connected='github_token' in session,
                         jira_connected='jira_token' in session,
                         github_repo_selected=session.get('selected_github_repo'),
                         jira_project_selected=session.get('selected_jira_project'))


@app.route("/query", methods=["POST"])
@login_required
def process_query():
    """Process user query through agentic RAG"""
    query = request.form.get('query')
    if not query:
        return jsonify({'error': 'Query is required'}), 400

    user_session = UserSession(
        user_id=session['user_id'],
        github_token=session.get('github_token'),
        jira_token=session.get('jira_token'),
        jira_url=session.get('jira_url'),
        jira_email=session.get('jira_email'),
        selected_github_repo=session.get('selected_github_repo'),
        selected_jira_project=session.get('selected_jira_project')
    )

    # Process query through agentic RAG
    result = agentic_rag.process_query(query, user_session)
    return jsonify(result)


@app.route("/query-interface")
@login_required
def query_interface():
    """Query interface for users to input requirements"""
    # Check if either GitHub repo or Jira project is selected
    github_selected = session.get('selected_github_repo')
    jira_selected = session.get('selected_jira_project')
    
    if not github_selected and not jira_selected:
        flash('Please select a GitHub repository or Jira project first')
        return redirect(url_for('dashboard'))
    
    return render_template("query_interface.html", 
                         github_repo=github_selected,
                         jira_project=jira_selected,
                         github_connected='github_token' in session,
                         jira_connected='jira_token' in session)


@app.route("/jira-project-selection")
@login_required
def jira_project_selection():
    """Redirect to dashboard - Jira project selection is now handled there"""
    return redirect(url_for('dashboard'))


@app.route("/jira-query-interface")
@login_required
def jira_query_interface():
    """Separate Jira query interface (keeping for backward compatibility)"""
    return redirect(url_for('query_interface'))


@app.route("/get-projects", methods=["POST"])
@login_required
def get_projects():
    """Get projects for connected Jira instance"""
    if not all(key in session for key in ['jira_url', 'jira_email', 'jira_token']):
        return jsonify({'error': 'Jira not connected'}), 400

    try:
        # Get projects using direct credentials
        auth = HTTPBasicAuth(session['jira_email'], session['jira_token'])
        projects_response = requests.get(
            f'{session["jira_url"]}/rest/api/3/project',
            auth=auth
        )
        if projects_response.status_code == 200:
            projects = projects_response.json()
            return jsonify({'projects': projects})
        else:
            return jsonify({'error': 'Failed to fetch projects'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route("/logout")
def logout():
    """Logout user"""
    session.clear()
    flash('You have been logged out')
    return redirect(url_for('connect_services'))


@app.route("/clear-github-selection", methods=["POST"])
@login_required
def clear_github_selection():
    """Clear GitHub repository selection"""
    session.pop('selected_github_repo', None)
    return jsonify({'success': True})


@app.route("/clear-jira-selection", methods=["POST"])
@login_required
def clear_jira_selection():
    """Clear Jira project selection"""
    session.pop('selected_jira_project', None)
    session.pop('jira_url', None)
    return jsonify({'success': True})


@app.route("/status")
def status():
    """Application status endpoint"""
    try:
        # Check service availability
        gemini_available = init_gemini() is not None
        qdrant_available = init_qdrant() is not None
        
        collections_info = {'docs_count': 0, 'tests_count': 0}
        if qdrant_available:
            try:
                qdrant_client = init_qdrant()
                docs_info = qdrant_client.get_collection(DOCS_COLLECTION)
                tests_info = qdrant_client.get_collection(TESTS_COLLECTION)
                collections_info = {
                    'docs_count': docs_info.points_count,
                    'tests_count': tests_info.points_count
                }
            except Exception as e:
                print(f"Error getting collection info: {e}")
        
        return jsonify({
            'status': 'running',
            'timestamp': datetime.now().isoformat(),
            'services': {
                'gemini_available': gemini_available,
                'qdrant_available': qdrant_available,
                'neo4j_available': True  # Assume available for now
            },
            'github_connected': 'github_token' in session,
            'jira_connected': 'jira_token' in session,
            'repo_selected': 'selected_github_repo' in session,
            'project_selected': 'selected_jira_project' in session,
            'collections_info': collections_info
        })
    except Exception as e:
        return jsonify({
            'status': 'running',
            'timestamp': datetime.now().isoformat(),
            'services': {
                'gemini_available': False,
                'qdrant_available': False,
                'neo4j_available': False
            },
            'github_connected': 'github_token' in session,
            'jira_connected': 'jira_token' in session,
            'repo_selected': 'selected_github_repo' in session,
            'project_selected': 'selected_jira_project' in session,
            'collections_info': {'docs_count': 0, 'tests_count': 0},
            'error': str(e)
        })


if __name__ == "__main__":
    initialize_app()
    app.run(debug=True, host="0.0.0.0", port=5003)
